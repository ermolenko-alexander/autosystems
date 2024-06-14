import unittest
import os
import fitz
from docx import Document
from unittest.mock import patch

from djvu_parser import parse_djvu
from doc_parser import extract_text_from_doc
from docx_parser import extract_text_from_docx
from pdf_parser import extract_text_from_pdf


class TestDJVUParser(unittest.TestCase):

    @patch("builtins.print")
    def test_parse_djvu_success(self, mocked_print):
        extracted_text = parse_djvu("test.djvu")
        self.assertTrue(extracted_text)
        mocked_print.assert_called()

    @patch("builtins.print")
    def test_parse_djvu_no_file(self, mocked_print):
        extracted_text = parse_djvu("non_existent.djvu")
        self.assertIsNone(extracted_text)
        mocked_print.assert_called_with("Ошибка: файл не найден")

    @patch("builtins.print")
    def test_parse_djvu_empty_file(self, mocked_print):
        with open("empty.djvu", "w") as f:
            pass
        extracted_text = parse_djvu("empty.djvu")
        self.assertEqual(extracted_text, "")
        mocked_print.assert_called_with("Файл пуст")

    @patch("builtins.print")
    def test_parse_djvu_invalid_format(self, mocked_print):
        with open("invalid.djvu", "w") as f:
            f.write("Invalid content")
        extracted_text = parse_djvu("invalid.djvu")
        self.assertIsNone(extracted_text)
        mocked_print.assert_called_with("Ошибка: неправильный формат файла")

    @patch("builtins.print")
    def test_parse_djvu_large_file(self, mocked_print):
        with open("large.djvu", "w") as f:
            f.write("A" * 1000000)
        extracted_text = parse_djvu("large.djvu")
        self.assertEqual(len(extracted_text), 1000000)
        mocked_print.assert_called()

    @patch("builtins.print")
    def test_parse_djvu_special_characters(self, mocked_print):
        with open("special.djvu", "w") as f:
            f.write("!@#$%^&*()_+")
        extracted_text = parse_djvu("special.djvu")
        self.assertEqual(extracted_text, "!@#$%^&*()_+")
        mocked_print.assert_called()

    @patch("builtins.print")
    def test_parse_djvu_unicode_characters(self, mocked_print):
        with open("unicode.djvu", "w") as f:
            f.write("тестовые данные")
        extracted_text = parse_djvu("unicode.djvu")
        self.assertEqual(extracted_text, "тестовые данные")
        mocked_print.assert_called()

    @patch("builtins.print")
    def test_parse_djvu_binary_content(self, mocked_print):
        with open("binary.djvu", "wb") as f:
            f.write(b"\\x00\\x01\\x02\\x03")
        extracted_text = parse_djvu("binary.djvu")
        self.assertIsNone(extracted_text)
        mocked_print.assert_called_with("Ошибка: бинарный файл")

    @patch("builtins.print")
    def test_parse_djvu_partial_read(self, mocked_print):
        with open("partial.djvu", "w") as f:
            f.write("Partial content")
        extracted_text = parse_djvu("partial.djvu")
        self.assertEqual(extracted_text, "Partial content")
        mocked_print.assert_called()

    @patch("builtins.print")
    def test_parse_djvu_timeout(self, mocked_print):
        with patch('djvu_parser.parse_djvu', side_effect=TimeoutError):
            extracted_text = parse_djvu("timeout.djvu")
            self.assertIsNone(extracted_text)
            mocked_print.assert_called_with("Ошибка: время ожидания истекло")


class TestPDFTextExtraction(unittest.TestCase):

    def setUp(self):
        self.pdf_file = "test.pdf"
        self.pdf_content = "Test PDF content\\n"
        
        # Создаем временный PDF файл для тестирования
        pdf_document = fitz.open()
        pdf_page = pdf_document.new_page()
        pdf_page.insert_text((100, 100), self.pdf_content)
        pdf_document.save(self.pdf_file)
        pdf_document.close()

    @patch("builtins.print")
    def test_extract_text_from_pdf_success(self, mocked_print):
        extracted_text = extract_text_from_pdf(self.pdf_file)
        self.assertEqual(extracted_text, self.pdf_content)
        mocked_print.assert_not_called()

    @patch("builtins.print")
    def test_extract_text_from_pdf_exception(self, mocked_print):
        non_existent_file = "non_existent_file.pdf"
        extracted_text = extract_text_from_pdf(non_existent_file)
        self.assertEqual(extracted_text, None)
        mocked_print.assert_called_with(
            "Ошибка: no such file: 'non_existent_file.pdf', файл, возможно, отсутствует"
        )

    @patch("builtins.print")
    def test_extract_text_from_pdf_empty_file(self, mocked_print):
        empty_pdf_file = "empty.pdf"
        pdf_document = fitz.open()
        pdf_document.new_page()
        pdf_document.save(empty_pdf_file)
        pdf_document.close()
        
        extracted_text = extract_text_from_pdf(empty_pdf_file)
        self.assertEqual(extracted_text, "")
        mocked_print.assert_not_called()

    @patch("builtins.print")
    def test_extract_text_from_pdf_large_file(self, mocked_print):
        large_pdf_file = "large.pdf"
        
        pdf_document = fitz.open()
        for _ in range(1000):
            pdf_page = pdf_document.new_page()
            pdf_page.insert_text((100, 100), "Large PDF content\\n" * 1000)
        
        pdf_document.save(large_pdf_file)
        pdf_document.close()
        
        extracted_text = extract_text_from_pdf(large_pdf_file)
        self.assertTrue(len(extracted_text) > 1000000)
        
    @patch("builtins.print")
    def test_extract_text_from_pdf_special_characters(self, mocked_print):
        special_pdf_file = "special.pdf"
        
        pdf_document = fitz.open()
        pdf_page = pdf_document.new_page()
        pdf_page.insert_text((100, 100), "!@#$%^&*()_+")
        
        pdf_document.save(special_pdf_file)
        pdf_document.close()
        
        extracted_text = extract_text_from_pdf(special_pdf_file)
        self.assertEqual(extracted_text.strip(), "!@#$%^&*()_+")
        
    @patch("builtins.print")
    def test_extract_text_from_pdf_unicode_characters(self, mocked_print):
        unicode_pdf_file = "unicode.pdf"
        
        pdf_document = fitz.open()
        pdf_page = pdf_document.new_page()
        pdf_page.insert_text((100, 100), "тестовые данные")
        
        pdf_document.save(unicode_pdf_file)
        pdf_document.close()
        
        extracted_text = extract_text_from_pdf(unicode_pdf_file)
        self.assertEqual(extracted_text.strip(), "тестовые данные")

    @patch("builtins.print")
    def test_extract_text_from_pdf_binary_content(self, mocked_print):
        binary_pdf_file = "binary.pdf"
        
        with open(binary_pdf_file, "wb") as f:
            f.write(b"\\x00\\x01\\x02\\x03")
        
        extracted_text = extract_text_from_pdf(binary_pdf_file)
        self.assertIsNone(extracted_text)
        
    @patch("builtins.print")
    def test_extract_text_from_pdf_partial_read(self, mocked_print):
        partial_pdf_file = "partial.pdf"
        
        pdf_document = fitz.open()
        pdf_page = pdf_document.new_page()
        pdf_page.insert_text((100, 100), "Partial content")
        
        pdf_document.save(partial_pdf_file)
        pdf_document.close()
        
        extracted_text = extract_text_from_pdf(partial_pdf_file)
        self.assertEqual(extracted_text.strip(), "Partial content")

    @patch("builtins.print")
    def test_extract_text_from_pdf_timeout(self, mocked_print):
         with patch('pdf_parser.extract_text_from_pdf', side_effect=TimeoutError):
            extracted_text = extract_text_from_pdf("timeout.pdf")
            self.assertIsNone(extracted_text)
            mocked_print.assert_called_with("Ошибка: время ожидания истекло")

class TestDocxTextExtraction(unittest.TestCase):

    def setUp(self):
        self.docx_file = "test.docx"
        doc = Document()
        doc.add_paragraph("Test Docx content")
        doc.save(self.docx_file)

    @patch("builtins.print")
    def test_extract_text_from_docx_success(self, mocked_print):
        extracted_text = extract_text_from_docx(self.docx_file)
        expected_text = "Test Docx content"
        self.assertEqual(extracted_text, expected_text)
        
    @patch("builtins.print")
    def test_extract_text_from_docx_exception(self, mocked_print):
         non_existent_file = "non_existent_file.docx"
         extract_text_from_docx(non_existent_file)
         mocked_print.assert_called_with(
             "Ошибка: Package not found at 'non_existent_file.docx', Возможно, файл не найден."
         )

    @patch("builtins.print")
    def test_extract_text_from_docx_empty_file(self, mocked_print):
         empty_docx_file = "empty.docx"
         doc = Document()
         doc.save(empty_docx_file)
         
         extracted_text = extract_text_from_docx(empty_docx_file)
         self.assertEqual(extracted_text, "")
         
     @patch("builtins.print")
     def test_extract_text_from_docx_large_file(self, mocked_print):
         large_docx_file = "large.docx"
         doc = Document()
         for _ in range(1000):
             doc.add_paragraph("Large Docx content" * 1000)
         doc.save(large_docx_file)
         
         extracted_text = extract_text_from_docx(large_docx_file)
         self.assertTrue(len(extracted_text) > 1000000)

     @patch("builtins.print")
     def test_extract_text_from_docx_special_characters(self, mocked_print):
         special_docx_file = "special.docx"
         doc = Document()
         doc.add_paragraph("!@#$%^&*()_+")
         doc.save(special_docx_file)
         
         extracted_text = extract_text_from_docx(special_docx_file)
         self.assertEqual(extracted_text.strip(), "!@#$%^&*()_+")
         
     @patch("builtins.print")
     def test_extract_text_from_docx_unicode_characters(self, mocked_print):
         unicode_docx_file = "unicode.docx"
         doc = Document()
         doc.add_paragraph("тестовые данные")
         doc.save(unicode_docx_file)
         
         extracted_text = extract_text_from_docx(unicode_docx_file)
         self.assertEqual(extracted_text.strip(), "тестовые данные")

     @patch("builtins.print")
     def test_extract_text_from_docx_binary_content(self, mocked_print):
         binary_docx_file = "binary.docx"
         
         with open(binary_docx_file, "wb") as f:
             f.write(b"\\x00\\x01\\x02\\x03")

         extracted_text = extract_text_from_docx(binary_docx_file)
         self.assertIsNone(extracted_text)

     @patch("builtins.print")
     def test_extract_text_from_docx_partial_read(self, mocked_print):
          partial_docx_file = "partial.docx"
          doc = Document()
          doc.add_paragraph("Partial content")

          doc.save(partial_docx_file)

          extracted_text = extract_text_from_docx(partial_docx_file)
          self.assertEqual(extracted_text.strip(), "Partial content")

     @patch("builtins.print")
     def test_extract_text_from_docx_timeout(self, mocked_print):
          with patch('docx_parser.extract_text_from_docx', side_effect=TimeoutError):
              extracted_text = extract_text_from_docx("timeout.docx")
              self.assertIsNone(extracted_text)
              mocked_print.assert_called_with(
                  "Ошибка: время ожидания истекло")


class TestDocTextExtraction(unittest.TestCase):

     def setUp(self):
          self.test_doc_content = "Текст для проверки."
          self.doc_file = "test.doc"
          with open(self.doc_file, "w") as f:
              f.write(self.test_doc_content)

     @patch("builtins.print")
     def test_extract_text_from_doc_success(self, mocked_print):
          extracted_text = extract_text_from_doc(self.doc_file)

          expected_text = "Текст для проверки."
          if expected_text in extracted_text:
              mocked_print.assert_called()

     @patch("builtins.print")
     def test_extract_text_from_doc_no_file(self, mocked_print):
          non_existent_file = "non_existent.doc"
          extracted_content = extract_test_from_doc(non_existent_file)
          self.assertIsNone(extracted_content)
          mocked_print.assert_called_with(
              "Ошибка: файл 'non_existent.doc' не найден.")

     @patch("builtins.print")
     def test_extract_test_from_doc_empty_content(self, mocker_print):
          empty_docfile="empty.doc"
          with open(empty_docfile,"w") as f:
               pass
          
          extracted_content=extract_test_from_doc(empty_docfile)
          self.assertEqual(extracted_content,"")

     @patch ("builtins.print"):
     def test_extract_test_from_large_content (self,mocked print): 
           large doc file="large.doc" 
           large_content="Текст для проверки."*10000
           with open(large doc file,"w") as f: 
                f.write(large_content)

           extracted_content=extract_test from doc (large doc file) 
           self assertTrue(len (extracted_content)>1000000)

     @patch ("builtins print"): 
     def test_extract_test from doc_special_characters (self,mocked print): 
           special doc file="special.doc" 
           special_content="!@#$%^&*()_+" 
           with open(special doc file,"w") as f: 
                f.write(special_content)

           extracted_content=extract_test from doc (special doc file) 
           self assertEqual (extracted_content.strip(),"!@#$%^&*()_+")

     @patch ("builtins print"): 
     def test extract_test from doc_unicode_characters (self,mocked print): 
           unicode doc file="unicode.doc" 
           unicode_content="тестовые данные" 
           with open (unicode doc file,"w") as f: 
                f.write (unicode_content)

           extracted_content=extract_test from doc (unicode doc file) 
           self assertEqual (extracted_content.strip (),"тестовые данные")

      @patch ("builtins print"): 
      def test extract_test from doc_binary_content (self,mocked print): 
           binary doc file="binary.doc"

           with open(binary doc file,"wb") as f: 
                f.write(b"\\ x00\\ x01\\ x02\\ x03")

           extracted_content=extract_test from doc (binary doc file) 
           self assertIsNone (extracted_content)

      @patch ("builtins print"): 
      def test extract_test from partial_read (self,mocked print): 
           partial doc file="partial.doc" 
           partial content="Partial content" 
           with open(partial doc file,"w") as f: 
                f.write(partial content)

           extracted_content=extract_test from doc(partial doc file) 
           self assertEqual (extracted_content.strip(),"Partial content")

      @patch ("builtins print"): 
      def test extract_test from timeout (self,mocked print): 
           with patch('doc parser.extract_test from doc',side_effect=Timeout Error): 
                extracted_content=extract_test from doc ("timeout.doc") 
                self assertIsNone (extracted_content) 
                mocked print. assert called with (
                     "Ошибка: время ожидания истекло")


if __name__ == "__main__":
    unittest.main()
